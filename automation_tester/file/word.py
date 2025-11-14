import io
import logging
from pathlib import Path

from docx import Document
from PIL import Image

from automation_tester.file.base_file import BaseFile
from automation_tester.utils.image_parser import parse_image_with_llm

logger = logging.getLogger(__name__)


class WordFile(BaseFile):
    """
    Word文档解析类，使用传统python-docx进行文本提取
    支持的格式：DOCX, DOC

    Args:
        source: 文件路径
    """

    def __init__(self, source: str):
        super().__init__(source)

    async def parse_text(self, **kwargs):
        if not self._path:
            raise ValueError("path is not set")

        file_path = Path(self._path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {self._path}")

        try:
            # 使用python-docx读取Word文档
            doc = Document(self._path)
            
            # 按顺序提取文档内容（文本、图片、表格）
            full_text = []
            await self._extract_content_in_order(doc, full_text)
            
            # 合并所有文本
            if full_text:
                combined_text = "\n".join(full_text)
                logger.info(f"解析结果: {combined_text}")
                if combined_text.strip():
                    yield combined_text.strip()
            else:
                logger.warning(f"未从Word文档中提取到文本: {self._path}")

        except Exception as e:
            logger.error(f"解析Word文档失败: {e}")
            raise RuntimeError(f"解析Word文档 {self._path} 失败: {e}")

    async def _extract_content_in_order(self, doc, full_text):
        try:
            # 创建图片映射，用于快速查找
            image_map = {}
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_map[rel.rId] = rel.target_part.blob
            
            # 遍历文档的所有元素，保持顺序
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    paragraph = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            paragraph = p
                            break
                    
                    if paragraph:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
                        
                        for run in paragraph.runs:
                            for inline_shape in run._element.xpath('.//a:blip'):
                                embed_id = inline_shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed_id and embed_id in image_map:
                                    image_data = image_map[embed_id]
                                    image_text = await self._parse_single_image(image_data)
                                    if image_text:
                                        full_text.append(f"{image_text}")
                
                elif element.tag.endswith('tbl'):
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    
                    if table:
                        table_text = []
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    table_text.append(cell.text.strip())
                        if table_text:
                            full_text.append(" | ".join(table_text))
        
        except Exception as e:
            logger.warning(f"按顺序提取内容时出错: {e}")
    async def _parse_single_image(self, image_data):
        """解析单个图片"""
        try:
            with Image.open(io.BytesIO(image_data)) as pil_image:
                pil_image.thumbnail((1024, 1024))
                return await parse_image_with_llm(pil_image)
        except Exception as e:
            logger.warning(f"解析单个图片时出错: {e}")
            return "图片解析失败"


